"""
tools/generate_docx.py — Gera arquivo .docx a partir dos dados do relatório.

Usa python-docx. Inclui cabeçalho institucional, título, tabela formatada
e rodapé com data de geração e protocolo.
"""
from __future__ import annotations

import io
import json
from datetime import datetime

import structlog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from langchain_core.tools import tool

from core.file_storage import save_file

logger = structlog.get_logger(__name__)

_PURPLE = RGBColor(0x43, 0x38, 0xCA)   # roxo Eleve
_GRAY   = RGBColor(0x55, 0x55, 0x55)


@tool
async def generate_docx(
    data: list,
    columns: list,
    title: str,
    school_id: str,
    school_name: str = "Escola",
    sa_token: str = "",
    **kwargs,
) -> str:
    """
    Gera relatório .docx e salva em file storage temporário.

    Args:
        data:        Lista de dicionários com os dados.
        columns:     Lista de colunas a exibir (chaves dos dicts).
        title:       Título do relatório.
        school_id:   ID da escola (usado no nome do arquivo).
        school_name: Nome da escola para o cabeçalho.

    Returns:
        JSON string com { file_id: str, rows: int }
    """
    try:
        doc = Document()

        # ── Margens ───────────────────────────────────────────────────
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1.2)
            section.right_margin  = Inches(1.2)

        # ── Cabeçalho ─────────────────────────────────────────────────
        school_para       = doc.add_paragraph()
        school_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = school_para.add_run(school_name)
        run.bold      = True
        run.font.size = Pt(14)
        run.font.color.rgb = _PURPLE

        title_para       = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.bold      = True
        run.font.size = Pt(12)

        date_para       = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(
            f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')} · Nicodemus ADM"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = _GRAY

        doc.add_paragraph()  # espaço

        # ── Tabela ────────────────────────────────────────────────────
        if data:
            table = doc.add_table(rows=1, cols=len(columns))
            table.style = "Table Grid"

            # Cabeçalhos
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(columns):
                hdr_cells[i].text = col_name.replace("_", " ").title()
                run = hdr_cells[i].paragraphs[0].runs[0]
                run.bold = True
                run.font.color.rgb = _PURPLE

            # Linhas de dados
            for row in data:
                row_cells = table.add_row().cells
                for i, col_name in enumerate(columns):
                    val = row.get(col_name)
                    if isinstance(val, list):
                        val = ", ".join(str(v) for v in val)
                    row_cells[i].text = str(val) if val is not None else "—"
        else:
            no_data = doc.add_paragraph("Nenhum registro encontrado para os filtros aplicados.")
            no_data.runs[0].font.color.rgb = _GRAY

        # ── Rodapé ────────────────────────────────────────────────────
        doc.add_paragraph()
        footer_para       = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer_para.add_run(
            f"Total de registros: {len(data)}"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = _GRAY

        # ── Serializa e salva ─────────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        file_id = save_file(school_id, buf.getvalue(), "docx")

        logger.info("generate_docx.ok", file_id=file_id, rows=len(data))
        return json.dumps({"file_id": file_id, "rows": len(data)})

    except Exception as exc:
        logger.error("generate_docx.error", error=str(exc))
        return json.dumps({"error": str(exc)})
